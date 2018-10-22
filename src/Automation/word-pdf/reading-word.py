from docx import Document
import os
# MAC 
# document = Document('/Users/samuelmensah/bulletin.docx')
# WINDOWS 
document = Document('bulletin2.docx')
for p in document.paragraphs:
    print (p.text)
